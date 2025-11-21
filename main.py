import os
from datetime import datetime, timedelta, timezone
from typing import Optional

from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer
from jose import JWTError, jwt
from passlib.context import CryptContext
from pydantic import BaseModel, EmailStr

from database import db, create_document, get_documents
from schemas import AppUser, DocxRequest

# JWT settings
SECRET_KEY = os.getenv("JWT_SECRET", "supersecretkeychange")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class Token(BaseModel):
    access_token: str
    token_type: str = "bearer"


class UserCreate(BaseModel):
    email: EmailStr
    password: str
    role: Optional[str] = "user"
    name: Optional[str] = None


class LoginRequest(BaseModel):
    email: EmailStr
    password: str


# Utility functions

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt


async def get_current_user(token: str = Depends(oauth2_scheme)) -> AppUser:
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception

    users = get_documents("appuser", {"email": email}, limit=1)
    if not users:
        raise credentials_exception

    doc = users[0]
    return AppUser(
        email=doc["email"],
        password_hash=doc["password_hash"],
        role=doc.get("role", "user"),
        name=doc.get("name"),
    )


@app.get("/")
def read_root():
    return {"message": "DOCX Formatter Backend Running"}


# Admin endpoint to create users
@app.post("/admin/create-user", response_model=dict)
async def admin_create_user(user: UserCreate):
    # Check if exists
    existing = get_documents("appuser", {"email": user.email}, limit=1)
    if existing:
        raise HTTPException(status_code=400, detail="User already exists")

    hashed = get_password_hash(user.password)
    data = AppUser(email=user.email, password_hash=hashed, role=user.role or "user", name=user.name)
    user_id = create_document("appuser", data)
    return {"id": user_id, "email": user.email}


# Login endpoint (JSON body)
@app.post("/auth/login", response_model=Token)
async def login(payload: LoginRequest):
    email = payload.email
    password = payload.password

    users = get_documents("appuser", {"email": email}, limit=1)
    if not users:
        raise HTTPException(status_code=400, detail="Incorrect email or password")

    user = users[0]
    if not verify_password(password, user["password_hash"]):
        raise HTTPException(status_code=400, detail="Incorrect email or password")

    access_token = create_access_token(data={"sub": email})
    return Token(access_token=access_token)


# DOCX generation
from io import BytesIO
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_LINE_SPACING


@app.post("/generate-docx")
async def generate_docx(payload: DocxRequest, user: AppUser = Depends(get_current_user)):
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Mm(payload.margin_top_mm)
    section.bottom_margin = Mm(payload.margin_bottom_mm)
    section.left_margin = Mm(payload.margin_left_mm)
    section.right_margin = Mm(payload.margin_right_mm)

    # Styles
    styles = doc.styles
    font_name = payload.font_family

    # Base (Normal)
    normal_style = styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(payload.font_size_paragraph or payload.font_size_global)

    # Heading styles
    for level, size in [(1, payload.h1_size), (2, payload.h2_size), (3, payload.h3_size)]:
        style = styles.get(f"Heading {level}")
        if style is None:
            style = styles.add_style(f"Heading {level}", 1)  # 1 = paragraph style
        style.font.name = font_name
        style.font.size = Pt(size)

    # Parse markdown-like text
    def add_paragraph(text_line: str):
        p = doc.add_paragraph(text_line)
        p_format = p.paragraph_format
        if payload.line_spacing:
            # Map common values to MS Word line spacing rules
            if abs(payload.line_spacing - 1.0) < 0.01:
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif abs(payload.line_spacing - 1.15) < 0.01:
                p_format.line_spacing = 1.15
            elif abs(payload.line_spacing - 1.5) < 0.01:
                p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif payload.line_spacing >= 2.0:
                p_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            else:
                p_format.line_spacing = payload.line_spacing
        return p

    lines = payload.text.splitlines()
    buffer = []

    def flush_paragraph():
        if buffer:
            add_paragraph(" ".join(buffer))
            buffer.clear()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(line[4:])
            p.style = styles["Heading 3"]
            continue
        if line.startswith("## "):
            flush_paragraph()
            p = doc.add_paragraph(line[3:])
            p.style = styles["Heading 2"]
            continue
        if line.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph(line[2:])
            p.style = styles["Heading 1"]
            continue
        buffer.append(line)

    flush_paragraph()

    # Output
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"formatted_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    headers = {"Content-Disposition": f"attachment; filename={filename}"}
    return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


@app.get("/test")
def test_database():
    """Test endpoint to check if database is available and accessible"""
    response = {
        "backend": "✅ Running",
        "database": "❌ Not Available",
        "database_url": None,
        "database_name": None,
        "connection_status": "Not Connected",
        "collections": []
    }

    try:
        if db is not None:
            response["database"] = "✅ Available"
            response["database_url"] = "✅ Configured"
            response["database_name"] = db.name if hasattr(db, 'name') else "✅ Connected"
            response["connection_status"] = "Connected"

            try:
                collections = db.list_collection_names()
                response["collections"] = collections[:10]
                response["database"] = "✅ Connected & Working"
            except Exception as e:
                response["database"] = f"⚠️  Connected but Error: {str(e)[:50]}"
        else:
            response["database"] = "⚠️  Available but not initialized"

    except Exception as e:
        response["database"] = f"❌ Error: {str(e)[:50]}"

    response["database_url"] = "✅ Set" if os.getenv("DATABASE_URL") else "❌ Not Set"
    response["database_name"] = "✅ Set" if os.getenv("DATABASE_NAME") else "❌ Not Set"

    return response


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
